from faster_whisper import WhisperModel
from typing import Tuple, Iterable
import imageio_ffmpeg as ffmpeg
from itertools import product
from docx import Document
from tqdm import tqdm
import pandas as pd
import warnings
import openpyxl
import requests
import difflib
import pickle
import shutil
import torch
import json
import csv
import os
import gc
import re

def flow_processing(csv_src_path, processed_csv_path, panel_master_path, intermediate_path, post_request_json):
    '''
    An automation script based on the Vs and Ps marks that you can use to transform the educational
    materials into JSON objects using the best AI world-wide techniques.

    -- csv_src_path        = the backend system mp3s. 
    -- processed_csv_path  = the processed CSV path.
    -- panel_master_path   = the main folder used as the center of the operations.
    -- intermediate_path   = the intermediate folder that will contain all the final JSONs from the content.
    -- post_request_json   = the final JSON file that is going directly to the backend system.

    '''
    
    def process_csv(csv_src_path, processed_csv_path, column_name = "Mp3", split_by = "/", replace_char = "_", new_column_name = "Course_Name"):
        '''
        Step 01: CSV File Editing & Panel Master Creating
        ــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        # Step 3: Read the CSV file
        df = pd.read_csv(csv_src_path)
    
        # Step 4: Extract information from the specified column
        def extract_info(url):
            parts = url.split(split_by)
            info = parts[-1].replace(replace_char, ' ').title()
            return info
    
        df[new_column_name] = df[column_name].apply(extract_info)
    
        # Step 5: Save the updated DataFrame to a new CSV file
        df.to_csv(processed_csv_path, index=False)
    
        print(f"DataFrame saved to '{processed_csv_path}'.")
    
        # Step 6: Remove the original CSV file
        if os.path.exists(csv_src_path):
            os.remove(csv_src_path)
            print(f"Original file removed.")
        else:
            print(f"Original file not found.")
    
    def create_course_folders(processed_csv_path, panel_master_path):
        '''
        Step 02: Creating Folders for Unique Course Names
        ــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        # Check if panel_master directory exists
        if not os.path.exists(panel_master_path):
            print("Error: panel_master directory not found.")
            return
    
        # Create a set to store unique course names
        unique_course_names = set()
    
        # Read the CSV file and extract course names
        with open(processed_csv_path, mode="r", encoding="utf-8") as csv_file:
            csv_reader = csv.DictReader(csv_file, delimiter=",")
            for row in csv_reader:
                course_name = row['Course_Name']
                unique_course_names.add(course_name)
    
        # Create folders for unique course names
        for course_name in unique_course_names:
            course_folder_path = os.path.join(panel_master_path, course_name)
            os.makedirs(course_folder_path, exist_ok=True)
            print(f"Created folder: {course_folder_path}")
    
    def download_mp3(mp3_url, panel_master_path, video_name, course_name, pbar):
        '''
        Step 03.01: Downloading the initial MP3 Files
        ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
    
        try:
            response = requests.get(mp3_url)
            response.raise_for_status()
    
            # Create the new file name based on the format: first 2 digits of "Video Name" - "Course_Name"
            new_mp3_name = f"{video_name[:2]}-{course_name}.mp3"
            mp3_drive_path = f"{panel_master_path}/{new_mp3_name}"
    
            with open(mp3_drive_path, "wb") as mp3_file:
                mp3_file.write(response.content)
    
            pbar.update(1)  # Update the collective progress bar
        except requests.RequestException as e:
            print(f"Failed to download {mp3_url}. Error: {e}")
    
    def download_and_rename_mp3(processed_csv_path, panel_master_path, mp3_column = "Mp3", course_column = "Course_Name", video_column = "Name"):
        '''
        Step 03.02: Downloading the processed MP3 Files and Renaming them
        ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        # Check if the panel_master directory exists
        if not os.path.exists(panel_master_path):
            print(f"Error: {panel_master_path} directory not found in Google Drive.")
            return
    
        # Read the CSV file
        with open(processed_csv_path, mode="r", encoding="utf-8") as csv_file:
            csv_reader = csv.DictReader(csv_file, delimiter=",")
    
            total_folders = len(os.listdir(panel_master_path))
            with tqdm(total=total_folders, desc="MP3 Downloading", unit="Audio File") as pbar:
                for row in csv_reader:
                    mp3_url = row[mp3_column]
                    course_name_csv = row[course_column]
                    video_name_csv = row[video_column]
    
                    # Check if the course_name in the CSV file matches any of the folders in panel_master
                    for folder_name in os.listdir(panel_master_path):
                        if course_name_csv.lower() == folder_name.lower():
                            full_panel_master_path = os.path.join(panel_master_path, folder_name)
    
                            # Download the mp3 file into the respective folder with the new name in Google Drive
                            download_mp3(mp3_url, full_panel_master_path, video_name_csv, course_name_csv, pbar)
    
    # change egyption context among the new prompts
    idx = 0
    initial_prompt_options = [
        'Transcribe this Egyptian speech into written text: هتشوف الحياة بطريقة مختلفة أوي عن الأول',
        'Transcribe this Egyptian speech into written text: طب لو أنا عايز أخس يبقى إيه هي المكملات اللي هتفيدني',
        ]
    
    def transcribe_mp3_files_faster_whisper(panel_master_path, processed_csv_path):
        '''
        Step 04: Debugging Mode for Transcription with Reference Control
        ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
    
        def is_sentence_matched(paragraph, sentence):
              # Tokenize the paragraph and sentence into words
              paragraph_tokens = paragraph.split()
              sentence_tokens = sentence.split()
    
              # Initialize a matcher using SequenceMatcher
              matcher = difflib.SequenceMatcher(None, paragraph_tokens, sentence_tokens)
    
              # Get matching blocks (sequences of matching words)
              matching_blocks = matcher.get_matching_blocks()
    
              # Calculate the total length of matching words
              total_matched_length = sum(match.size for match in matching_blocks)
    
              # Determine if the sentence is matched
              sentence_length = len(sentence_tokens)
              return total_matched_length / sentence_length
    
        file_path = "/content/drive/MyDrive/Final_Automation/egy_reference/egy_reference.txt"
        with open(file_path, 'r', encoding='utf-8') as file:
            paragraph = file.read().replace("\n", ' ')
    
        global idx
        # Load CSV file to get the total file count
        df = pd.read_csv(processed_csv_path)
        if 'Course_Name' not in df.columns:
            print("Error: 'Course_Name' column missing in CSV!")
            return
        total_files = int(df['Course_Name'].nunique())
        print(f"Total unique courses in CSV: {total_files}")
    
        def has_transcription_json(folder_path):
            for filename in os.listdir(folder_path):
                if filename.endswith('Transcriptions.json'):
                    return True
            return False
    
        for folder_name in os.listdir(panel_master_path):
            full_folder_drive_path = os.path.join(panel_master_path, folder_name)
    
            if full_folder_drive_path.endswith('.docx') or full_folder_drive_path.endswith('.txt') or full_folder_drive_path.endswith('.csv') or full_folder_drive_path.endswith(".ipynb_checkpoints"):
                continue
    
            mp3_files = [f for f in os.listdir(full_folder_drive_path) if f.lower().endswith('.mp3')]
            if len(mp3_files) == 0:
                print(f"No MP3 files found in {folder_name}")
                continue
    
            mp3_files.sort(key=lambda x: int(x.split("-")[0].strip()) if x.split("-")[0].strip().isdigit() else float('inf'))
            print(f"Found {len(mp3_files)} MP3 files in {folder_name}")
    
            # Initialize the model and load weights before processing MP3 files
            model_size = "large-v2"
            model = WhisperModel(model_size, device="cuda", compute_type="float16")
    
            mp3_files_progress = tqdm(mp3_files, desc=f"\nTranscribing {folder_name}", unit="File")
            course_transcription = {}
    
            for file_name in mp3_files_progress:
                full_file_drive_path = os.path.join(full_folder_drive_path, file_name)
    
                try:
                    sql_id = int(file_name[:2])
                except ValueError:
                    print(f"Error extracting SQL ID from {file_name}")
                    continue
                print(f"Extracted SQL ID {sql_id} from {file_name}")
    
                course_name_parts = folder_name.split("-")
                course_name = course_name_parts[1].strip() if len(course_name_parts) > 1 else folder_name.strip()
    
                matching_rows = df[(df['Course_Name'] == course_name) & (df['Name'].str[:2] == str(sql_id).zfill(2))]
                if not matching_rows.empty:
                    video_id = int(matching_rows['Id'].values[0])
                    print(f"Matched video ID {video_id} for SQL ID {sql_id} in {course_name}")
                else:
                    video_id = None
                    print(f"No matching video ID found for SQL ID {sql_id} in {course_name}")
                flag = True
                while flag:
                    segments_g, _ = model.transcribe(full_file_drive_path,
                                                  vad_filter=True,
                                                  beam_size = 11,
                                                  best_of = 9,
                                                  word_timestamps = True,
                                                  no_speech_threshold = 0.2,
                                                  vad_parameters = dict(min_silence_duration_ms = 2000),
                                                  initial_prompt = initial_prompt_options[idx % len(initial_prompt_options)]
                                                  )
    
                    segments = []
                    segment_id = 0
                    for segment in segments_g:
                        segments.append({})
                        segments[segment_id]['start'] = segment.start
                        segments[segment_id]['end'] = segment.end
                        segments[segment_id]['text'] = segment.text
                        segment_id += 1
    
                    ## if avg prob of segment matched egyption more than 50% will sucessed.
                    summation_prob = 0
                    for segment in segments:
                        summation_prob += is_sentence_matched(paragraph, segment['text'])
                    avg_prob = summation_prob / len(segments)
    
                    if avg_prob >= 0.6:
                        print(f"\nSuccessful initial sentence: {segments[0]['text']} of video_{sql_id} has avg_prob = {int(avg_prob*100)}%")
                        print(f"The initial prompt is: [ {initial_prompt_options[idx % len(initial_prompt_options)]} ]")
                        flag = False
                    else:
                        print(f"FAILED at the initial sentence {segments[0]['text']} with avg_prob = {int(avg_prob * 100)}%")
                        print(f"The initial prompt is: [ {initial_prompt_options[idx % len(initial_prompt_options)]} ]")
    
                        # Condition override technique
                        user_input = input(f"\nDo you want to override and accept this as 0.6 matching? (y/n): ")
                        if user_input.lower() == 'y':
                            flag = False
    
                    idx += 1
    
                if not segments:
                    print(f"No segments/transcriptions found for {file_name}")
                    continue
    
                entry_list = []
                view_index = 1
                for segment in segments:
                    entry = {
                        "videoId": video_id,
                        "sqlId": sql_id,
                        "paragraphInfo": {
                            "id": "",
                            "viewIndex": view_index,
                            "startWord": "",
                            "endWord": "",
                            "startSecond": int(segment['start']),
                            "endSecond": int(segment['end']),
                            "paragraphDetails": segment['text'],
                            "objectiveId": "",
                            "skillsInfo": [
                                {
                                    "skillId": ""
                                }
                            ]
                        }
                    }
                    entry_list.append(entry)
                    view_index += 1
    
                if sql_id not in course_transcription:
                    course_transcription[sql_id] = []
                course_transcription[sql_id].extend(entry_list)
    
            mp3_files_progress.close()
    
            # Release model memory and clear GPU cache
            gc.collect()
            del model
            torch.cuda.empty_cache()
    
            if not course_transcription:
                print(f"No transcriptions generated for {folder_name}")
                continue
    
            merged_doc = Document()
            for sql_id, entries in course_transcription.items():
                video_id_str = f"[V{sql_id}]"
                merged_doc.add_paragraph(video_id_str, style='Normal')
                for entry in entries:
                    timestamp = f"Start: {entry['paragraphInfo']['startSecond']:.2f}s,  End: {entry['paragraphInfo']['endSecond']:.2f}s"
                    merged_doc.add_paragraph(timestamp, style='Normal')
                    merged_doc.add_paragraph(entry['paragraphInfo']['paragraphDetails'])
    
            merged_docx_drive_path = os.path.join(full_folder_drive_path, f"{folder_name} AltTranscriptions.docx")
            merged_doc.save(merged_docx_drive_path)
            print(f"Saved transcriptions for {folder_name} to {merged_docx_drive_path}")
    
            course_json_drive_path = os.path.join(full_folder_drive_path, f"{folder_name} Transcriptions.json")
            with open(course_json_drive_path, "w", encoding="utf-8") as json_file:
                json.dump(course_transcription, json_file, ensure_ascii=False, indent=4)
            print(f"Saved JSON for {folder_name} to {course_json_drive_path}")
    
            if has_transcription_json(full_folder_drive_path):
                lectures_folder_path = os.path.join(full_folder_drive_path, 'lectures_folder')
                if not os.path.exists(lectures_folder_path):
                    os.mkdir(lectures_folder_path)
                for filename in os.listdir(full_folder_drive_path):
                    file_path = os.path.join(full_folder_drive_path, filename)
                    if filename.endswith('.mp3'):
                        shutil.move(file_path, os.path.join(lectures_folder_path, filename))
    
        print(f"Transcriptions for all courses completed successfully")
    
    def move_files_to_folders(content_directory, panel_master_path, file_extensions = ['.xlsx', '.docx']):
        '''
        Step 05: Looping on content_files and restructure panel_master
        ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        # Get a list of files in the source directory
        files_to_move = os.listdir(content_directory)
    
        # Loop through files and move them
        for filename in files_to_move:
            # Check if the file has one of the specified extensions
            if any(filename.endswith(ext) for ext in file_extensions):
                # Extract the base file name (without extension)
                file_base_name = os.path.splitext(filename)[0]
    
                # Create the expected folder name in the destination directory
                expected_folder_name = os.path.join(panel_master_path, file_base_name)
    
                # Check if the folder with the same name exists in the destination directory
                if os.path.exists(expected_folder_name):
                    # Move the file to the corresponding folder
                    shutil.move(os.path.join(content_directory, filename), expected_folder_name)
                else:
                    # If the folder doesn't exist, create it and then move the file
                    os.makedirs(expected_folder_name)
                    shutil.move(os.path.join(content_directory, filename), expected_folder_name)
    
        # Provide a summary and handle any potential errors
        print("File move operation completed successfully!")
    
    def pyillam_script_final(docx_file_path):
        '''
        Step 06.01: Document Cleaning & Paragraphs Extraction
        ــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        document = Document(docx_file_path)
        text = "\n".join([paragraph.text for paragraph in document.paragraphs]).strip().lower()
    
        lines = text.split('\n')
        final_lines = [i.strip() for i in lines if len(i) > 1]
    
        v_list = []
        video_info = []
        titles = []
        for i, element in enumerate(final_lines):
            if re.search(r'v\d', element):
                v_list.append(i)
    
        for video_index in v_list:
            v_line = final_lines[video_index].strip()
            v_index = v_line[v_line.find('v')]
            video_title = final_lines[video_index + 1].strip().title()
            titles.append(video_title)
            video_info.append({'video_index': v_index, 'video_title': video_title})
    
        script = [title for title in final_lines if title.lower() not in [t.lower() for t in titles]]
    
        v_trimmed_list = [j for j, j_element in enumerate(script) if re.search(r'v\d', j_element)]
        v_trimmed_list.append(len(script))
    
        paragraphs = [script[start_idx:end_idx][1:] for start_idx, end_idx in zip(v_trimmed_list[:-1], v_trimmed_list[1:])]
    
        index_lists = [[int(re.search(r'p(\d+)', element).group(1)) for element in sub_list if re.search(r'p\d', element)] for sub_list in paragraphs]
        videos_list = [[jendex for jendex, element in enumerate(sub_list) if re.search(r'p\d', element)] for sub_list in paragraphs]
    
        for sub_list in paragraphs:
            sub_list.append(len(sub_list))
    
        paragraphs_result = []
        for i in range(len(index_lists)):
            paragraphs_result.append({'videoId': i + 1, 'video_title': video_info[i]['video_title'], 'paragraphInfo': []})
            paragraphs_counter = 1
            for j in range(len(videos_list[i])):
                start = videos_list[i][j] + 1
                end = videos_list[i][j + 1] if j + 1 < len(videos_list[i]) else -1
                paragraphs_result[i]['paragraphInfo'].append({'viewIndex': paragraphs_counter, 'paragraphDetails': " ".join(paragraphs[i][start:end])})
                paragraphs_counter += 1
    
        return paragraphs_result
    
    def process_docx_files_final(panel_master_path):
        '''
        Step 06.02: Document Analysis Phase
        ــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        for folder_name in os.listdir(panel_master_path):
            full_folder_drive_path = os.path.join(panel_master_path, folder_name)
    
            # Check if the folder contains a .docx file with the folder's name
            docx_file_path = os.path.join(full_folder_drive_path, f"{folder_name}.docx")
            if os.path.isfile(docx_file_path):
                script_result = pyillam_script_final(docx_file_path)
    
                # Save as a JSON file
                json_file_path = os.path.join(full_folder_drive_path, f"{folder_name} Script.json")
                with open(json_file_path, "w", encoding="utf-8") as json_file:
                    json.dump(script_result, json_file, ensure_ascii=False, indent=4)
    
    def extract_number(text):
        match = re.search(r'\d+', str(text))
        return int(match.group()) if match else None
    
    def is_numeric(text):
        try:
            float(text)
            return True
        except ValueError:
            return False
    
    def get_correct_choices(right_answer, num_choices):
        correct_choices = []
        if not pd.isna(right_answer) and isinstance(right_answer, str):
            try:
                digits = [int(digit) for digit in re.findall(r'\d', right_answer)]
                for digit in digits:
                    if 1 <= digit <= num_choices:
                        correct_choices.append(digit)
            except ValueError:
                pass
        return correct_choices
    
    def process_excel_files(panel_master_path):
        '''
        Step 07: Excel File Analysis - Part One
        ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
    
        for folder in os.listdir(panel_master_path):
            full_folder_drive_path = os.path.join(panel_master_path, folder)
    
            if os.path.isdir(full_folder_drive_path):
                for file in os.listdir(full_folder_drive_path):
                    if file.endswith('.xlsx'):
                        full_xlsx_path = os.path.join(full_folder_drive_path, file)
                        with warnings.catch_warnings():
                            warnings.filterwarnings("ignore", category=UserWarning, message="Unknown extension is not supported and will be removed")
                            df = pd.read_excel(full_xlsx_path)
    
                        df.columns = df.columns.str.lower()
    
                        def pyillam_excelinfo():
                            questions_info = []
                            for index, row in df.iterrows():
                                # alternative = int(row['alt']) if not pd.isna(row['alt']) else None
                                questionTypeId = str(row['question type'])
    
                                if questionTypeId.lower() == "true&false":
                                    questionTypeId = 1
                                else:
                                    questionTypeId = 2
    
                                question_info = {
                                    "id": "",
                                    "questionTypeId": questionTypeId,
                                    "question_videoId": extract_number(row['v']),
                                    "question_paragraph": extract_number(row['p']),
                                    "questionDetails": None if pd.isna(row.get('question', None)) else row['question'],
                                    "timeLimit": 45,
                                    "skipping": False,
                                    "preAssessment": is_numeric(row['c']),
                                    "chapter": extract_number(row['c']) if is_numeric(row['c']) else None,
                                    "finalExam": not is_numeric(row['c']),
                                    "pathways": True,
                                    "games": True,
                                    "alternative": 0,
                                    "questions_skills_objectives": [
                                        {
                                            "skill_Id": extract_number(row['s']),
                                            "objective_Id": extract_number(row['l']),
                                            "level_Id": ""
                                        }
                                    ],
                                    "questionAnswers": []
                                }
    
                                total_points = 10 if question_info['questions_skills_objectives'][0]['level_Id'] in [1, 2] else 20 if question_info['questions_skills_objectives'][0]['level_Id'] in [3, 4] else 30
    
                                correct_choices = get_correct_choices(row['right answer'], 6)
    
                                if not correct_choices:
                                    print(f"Empty list found in {file} at row {index + 2}")
    
                                correct_choices_num = len(correct_choices)
    
                                if str(row['question type']).lower() == "select":
                                    for i in range(1, 7):
                                        choice = row.get(f'choice {i}', None)
                                        answers = None if pd.isna(choice) else choice
                                        point = total_points // correct_choices_num if i in correct_choices else 0
    
                                        if i in correct_choices:
                                            correctAnswers = answers
                                        else:
                                            correct_answer_index = correct_choices[0]
                                            correctAnswers = row.get(f'choice {correct_answer_index}', None)
    
                                        question_info["questionAnswers"].append({
                                            "viewIndex": i,
                                            "answers": str(answers),
                                            "correctAnswers": str(correctAnswers),
                                            "point": point
                                        })
    
                                elif str(row['question type']).lower() == "mcq":
                                    for i in range(1, 5):
                                        choice = row.get(f'choice {i}', None)
                                        answers = None if pd.isna(choice) else choice
                                        point = total_points // correct_choices_num if i in correct_choices else 0
    
                                        if i in correct_choices:
                                            correctAnswers = answers
                                        else:
                                            correct_answer_index = correct_choices[0]
                                            correctAnswers = row.get(f'choice {correct_answer_index}', None)
    
                                        question_info["questionAnswers"].append({
                                            "viewIndex": i,
                                            "answers": str(answers),
                                            "correctAnswers": str(correctAnswers),
                                            "point": point
                                        })
    
                                else:
                                    for i in range(1, 3):
                                        choice = row.get(f'choice {i}', None)
                                        if isinstance(choice, bool):
                                            answers = "True" if choice else "False"
                                        else:
                                            answers = "True" if choice and choice == "True" else "False"
                                        point = total_points if i in correct_choices else 0
    
                                        if i in correct_choices:
                                            correctAnswers = str(answers)
                                        else:
                                            correct_answer_index = correct_choices[0]
                                            correctAnswers = row.get(f'choice {correct_answer_index}', None)
    
                                        if correctAnswers == "true":
                                            correctAnswers == "True"
                                        elif correctAnswers == "":
                                            correctAnswers == "True"
    
                                        question_info["questionAnswers"].append({
                                            "viewIndex": i,
                                            "answers": str(answers),
                                            "correctAnswers": str(correctAnswers),
                                            "point": point
                                        })
    
                                questions_info.append(question_info)
    
                            course_title = df.iloc[0]['course name']
                            json_structure = {
                                "course title": course_title,
                                "questionsInfo": questions_info
                            }
    
                            questions_info_json = json.dumps(json_structure, indent=4, ensure_ascii=False)
                            return questions_info_json
    
                        json_data = pyillam_excelinfo()
    
                        if json_data:
                            folder_name = os.path.basename(full_folder_drive_path)
                            json_file_name = f"{folder_name} Quiz.json"
                            json_file_path = os.path.join(full_folder_drive_path, json_file_name)
                            with open(json_file_path, 'w', encoding='utf-8') as json_file:
                                json_file.write(json_data)
    
    def extract_skills_objectives(panel_master_path):
        '''
        Step 08: Excel File Analysis - Part Two
        ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        for folder in os.listdir(panel_master_path):
            full_folder_drive_path = os.path.join(panel_master_path, folder)
    
            if os.path.isdir(full_folder_drive_path):
                map_objective_item = {}
                map_S_item = {}
    
                for file in os.listdir(full_folder_drive_path):
                    if file.endswith('.xlsx'):
                        full_xlsx_path = os.path.join(full_folder_drive_path, file)
                        df = pd.read_excel(full_xlsx_path, sheet_name=1)
                        df.columns = df.columns.str.lower()
    
                        for index, row in df.iterrows():
                            if pd.notna(row['index']):
                                current_index = row['index']
                                if any(char.lower() == 'l' for char in current_index) and any(char.isdigit() for char in current_index):
                                    digit_index = ''.join(filter(str.isdigit, current_index))
                                    if digit_index not in map_objective_item:
                                        map_objective_item[digit_index] = str(row['item (english)']).strip()
                                    else:
                                        map_objective_item[digit_index] += ', ' + str(row['item (english)'])
                                elif any(char.lower() == 's' for char in current_index) and any(char.isdigit() for char in current_index):
                                    digit_index = ''.join(filter(str.isdigit, current_index))
                                    if digit_index not in map_S_item:
                                        map_S_item[digit_index] = str(row['item (english)']).strip()
                                    else:
                                        map_S_item[digit_index] += ', ' + str(row['item (english)'])
                                else:
                                    pass
    
                skills_json_filename = os.path.join(full_folder_drive_path, f"{folder} Skills.json")
                objectives_json_filename = os.path.join(full_folder_drive_path, f"{folder} Objectives.json")
    
                with open(skills_json_filename, 'w') as skills_json_file:
                    json.dump(map_S_item, skills_json_file, indent=4)
    
                with open(objectives_json_filename, 'w') as objectives_json_file:
                    json.dump(map_objective_item, objectives_json_file, indent=4)
    
    def update_questions_with_skills_objectives(panel_master_path):
        '''
        Step 09: Excel File Analysis - Part Three
        ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        for root, _, files in os.walk(panel_master_path):
            for file in files:
                if file.endswith('Quiz.json'):
                    quiz_file_path = os.path.join(root, file)
    
                    folder_name = os.path.basename(root)
    
                    skill_id_mapping = {}
                    objective_id_mapping = {}
    
                    skills_file_path = os.path.join(root, folder_name + ' Skills.json')
                    objectives_file_path = os.path.join(root, folder_name + ' Objectives.json')
    
                    if os.path.exists(skills_file_path):
                        with open(skills_file_path, 'r', encoding='utf-8') as skills_file:
                            skills_data = json.load(skills_file)
    
                        skill_id_mapping.update({
                            str(skill_id): skill_name for skill_id, skill_name in skills_data.items()
                        })
    
                    if os.path.exists(objectives_file_path):
                        with open(objectives_file_path, 'r', encoding='utf-8') as objectives_file:
                            objectives_data = json.load(objectives_file)
    
                        objective_id_mapping.update({
                            str(objective_id): objective_name for objective_id, objective_name in objectives_data.items()
                        })
    
                    with open(quiz_file_path, 'r', encoding='utf-8') as quiz_file:
                        quiz_data = json.load(quiz_file)
    
                    for question in quiz_data['questionsInfo']:
                        skill_id = str(question['questions_skills_objectives'][0]['skill_Id'])
                        if skill_id in skill_id_mapping:
                            question['questions_skills_objectives'][0]['skill_Id'] = skill_id_mapping[skill_id]
    
                        objective_id = str(question['questions_skills_objectives'][0]['objective_Id'])
                        if objective_id in objective_id_mapping:
                            question['questions_skills_objectives'][0]['objective_Id'] = objective_id_mapping[objective_id]
    
                    output_file_path = os.path.join(root, f'{folder_name} Updated Questions.json')
    
                    with open(output_file_path, 'w', encoding='utf-8') as updated_quiz_file:
                        json.dump(quiz_data, updated_quiz_file, ensure_ascii=False, indent=4)
    
                    print(f"Updated Questions JSON for {folder_name} saved to {output_file_path}")
    
                    combined_mappings = {
                        'course_skills': skill_id_mapping,
                        'course_objectives': objective_id_mapping
                    }
    
                    combined_mappings_file_path = os.path.join(root, f'{folder_name} Skills&Objectives.json')
    
                    with open(combined_mappings_file_path, 'w', encoding='utf-8') as combined_mappings_file:
                        json.dump(combined_mappings, combined_mappings_file, ensure_ascii=False, indent=4)
    
                    print(f"Combined mappings for {folder_name} saved to {combined_mappings_file_path}")
    
        print("Processing completed.")
    
    def transform_data_to_desired_format(script_data, transcriptions_data, questions_data, sim_percentage = 0.7):
        '''
        Step 10.01: Final NLP Matching for transforming data into the desired format
        ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
        
        def is_sentence_matched(paragraph, sentence):
            # Tokenize the paragraph and sentence into words
            paragraph_tokens = paragraph.split()
            sentence_tokens = sentence.split()
    
            # Initialize a matcher using SequenceMatcher
            matcher = difflib.SequenceMatcher(None, paragraph_tokens, sentence_tokens)
    
            # Get matching b|locks (sequences of matching words)
            matching_blocks = matcher.get_matching_blocks()
    
            # Calculate the total length of matching words
            total_matched_length = sum(match.size for match in matching_blocks)
    
            # Determine if the sentence is matched (you can set a threshold as needed)
            sentence_length = len(sentence_tokens)
            return total_matched_length / sentence_length
    
        final_result = []
    
        for video_key, transcriptions in transcriptions_data.items():
            video_title = next((video['video_title'] for video in script_data if video['videoId'] == int(video_key)), 'Default Video Title')
            video_data = {
                "videoId": transcriptions[0]['videoId'],
                "video_title": video_title,
                "sqlId": 1,
                "paragraphInfo": []
            }
    
            for paragraph_info in script_data:
                if paragraph_info['videoId'] == int(video_key):
                    for paragraph_data in paragraph_info['paragraphInfo']:
                        updated_paragraph_data = {
                            "id": "",
                            "viewIndex": paragraph_data['viewIndex'],
                            "startWord": " ".join(paragraph_data['paragraphDetails'].split()[:2]),
                            "endWord": " ".join(paragraph_data['paragraphDetails'].split()[-2:]),
                            "startSecond": None,
                            "endSecond": None,
                            "paragraphDetails": paragraph_data['paragraphDetails'],
                            "objectiveId": "",
                            "skillsInfo": [{"skillId": ""}],
                            "questionsInfo": []
                        }
    
                        # print("Actual Paragraph:", paragraph_data['paragraphDetails'])
                        for transcription in transcriptions:
                            sentence = transcription['paragraphInfo']['paragraphDetails']
                            l = len(sentence)
                            if is_sentence_matched(paragraph_data['paragraphDetails'], sentence) >= sim_percentage:
                                start_second = transcription['paragraphInfo']['startSecond']
                                end_second = transcription['paragraphInfo']['endSecond']
    
                                if updated_paragraph_data['startSecond'] is None or updated_paragraph_data['startSecond'] > start_second:
                                    updated_paragraph_data['startSecond'] = start_second
    
                                if updated_paragraph_data['endSecond'] is None or updated_paragraph_data['endSecond'] < end_second:
                                    updated_paragraph_data['endSecond'] = end_second
    
                        for question_info in questions_data['questionsInfo']:
                            if question_info['question_videoId'] == paragraph_info['videoId'] and question_info['question_paragraph'] == paragraph_data['viewIndex']:
                                skills_objectives = question_info.get('questions_skills_objectives', [{}])[0]
                                updated_paragraph_data['objectiveId'] = str(skills_objectives.get('objective_Id', None))
                                updated_paragraph_data['skillsInfo'][0]['skillId'] = skills_objectives.get('skill_Id', None)
    
                                updated_question_info = {
                                    "id": "",
                                    "questionTypeId": question_info['questionTypeId'],
                                    "questionDetails": question_info['questionDetails'],
                                    "timeLimit": question_info['timeLimit'],
                                    "skipping": question_info['skipping'],
                                    "preAssessment": question_info['preAssessment'],
                                    "chapter": not question_info['finalExam'],
                                    "finalExam": question_info['finalExam'],
                                    "pathways": question_info['pathways'],
                                    "games": question_info['games'],
                                    "alternative": question_info['alternative'],
                                    "questionsSkills": [{"skillId": str(skills_objectives.get('skill_Id', None))}],
                                    "questionAnswers": question_info['questionAnswers']
                                }
                                updated_paragraph_data['questionsInfo'].append(updated_question_info)
    
                        video_data['paragraphInfo'].append(updated_paragraph_data)
            final_result.append(video_data)
        return final_result
    
    def process_subfolder(subfolder_path):
        # Get the folder name from the subfolder path
        folder_name = os.path.basename(subfolder_path)
    
        # Check for JSON files with specific suffixes in the subfolder
        script_file_path = os.path.join(subfolder_path, f"{folder_name} Script.json")
        transcriptions_file_path = os.path.join(subfolder_path, f"{folder_name} Transcriptions.json")
        questions_file_path = os.path.join(subfolder_path, f"{folder_name} Updated Questions.json")
    
        # Check if all required JSON files exist
        if os.path.exists(script_file_path) and os.path.exists(transcriptions_file_path) and os.path.exists(questions_file_path):
            # Load the JSON files
            with open(script_file_path, "r") as script_file:
                script_data = json.load(script_file)
    
            with open(transcriptions_file_path, "r") as transcriptions_file:
                transcriptions_data = json.load(transcriptions_file)
    
            with open(questions_file_path, "r") as questions_file:
                questions_data = json.load(questions_file)
    
            # Transform the data using the new logic
            transformed_data = transform_data_to_desired_format(script_data, transcriptions_data, questions_data)
    
            # Define the output file path
            output_file_path = os.path.join(subfolder_path, f"{folder_name} Final.json")
    
            # Save the transformed data to a JSON file
            with open(output_file_path, "w", encoding="utf-8") as output_file:
                json.dump(transformed_data, output_file, ensure_ascii=False, indent=4)
    
            print(f"Transformed data saved to: {output_file_path}")
        else:
            print(f"Required JSON files not found in subfolder: {subfolder_path}")
    
    def final_matching(panel_master_path):
        '''
        Step 10.02: Final NLP Matching for difflip
        ــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
    
        # Loop through all sub-folders in the root directory
        for subfolder_name in os.listdir(panel_master_path):
            subfolder_path = os.path.join(panel_master_path, subfolder_name)
    
            # Check if the item is a directory
            if os.path.isdir(subfolder_path):
                process_subfolder(subfolder_path)
    
    def format_paragraph_info(paragraph):
        '''
        Step 11: Final NLP Matching for reformatting the JSON output
        ــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ
        '''
    
        default_paragraph = {
            "id": "",
            "viewIndex": 0,
            "startWord": "",
            "endWord": "",
            "startSecond": 0,
            "endSecond": 0,
            "paragraphDetails": "",
            "objectiveId": "",
            "skillsInfo": [],
            "questionsInfo": []
        }
        formatted_paragraph = {key: paragraph.get(key, default_paragraph[key]) for key in default_paragraph}
        formatted_paragraph["startSecond"] = paragraph.get("startSecond", 0) or 0
        formatted_paragraph["endSecond"] = paragraph.get("endSecond", 0) or 0
        formatted_paragraph["skillsInfo"] = [skill for skill in paragraph.get("skillsInfo", []) if skill.get("skillId")]
        formatted_paragraph["questionsInfo"] = [format_question_info(question) for question in paragraph.get("questionsInfo", [])]
        return formatted_paragraph
    
    def format_question_info(question):
        default_question = {
            "id": "",
            "questionTypeId": 0,
            "questionDetails": "",
            "timeLimit": 0,
            "skipping": True,
            "preAssessment": True,
            "chapter": True,
            "finalExam": True,
            "pathways": True,
            "games": True,
            "alternative": 0,
            "questionsSkills": [],
            "questionAnswers": []
        }
        formatted_question = {key: question.get(key, default_question[key]) for key in default_question}
        formatted_question["questionAnswers"] = [
            {
                "viewIndex": ans.get("viewIndex", 0),
                "answers": ans.get("answers", "") or "",
                "correctAnswers": ans.get("correctAnswers", "") or "",
                "point": ans.get("point", 0) or 0
            }
            for ans in question.get("questionAnswers", [])
        ]
        return formatted_question
    
    def transform_json_content(json_content):
        formatted_content = []
        for video in json_content:
            formatted_video = {
                "videoId": video["videoId"],
                "video_title": video["video_title"],
                "paragraphInfo": [format_paragraph_info(paragraph) for paragraph in video.get("paragraphInfo", [])]
            }
            formatted_content.append(formatted_video)
        filtered_content = [video for video in formatted_content if video["paragraphInfo"]]
        return {"videosScriptsInfo": filtered_content}
    
    def process_json_file(json_file_path):
        with open(json_file_path, "r", encoding='utf-8') as file:
            json_content = json.load(file)
            transformed_content = transform_json_content(json_content)
    
        with open(json_file_path, "w", encoding='utf-8') as output_file:
            json.dump(transformed_content, output_file, ensure_ascii=False, indent=4)
    
    def process_all_json_files_in_folder(panel_master_path):
        for root, dirs, files in os.walk(panel_master_path):
            for file in files:
                if file.endswith("Final.json"):
                    json_file_path = os.path.join(root, file)
                    process_json_file(json_file_path)
    
    def merge_jsons(panel_master_path, intermediate_path, post_request_json):
        # Copy JSON files from subfolders to destination folder
        for folder_name in os.listdir(panel_master_path):
            full_folder_path = os.path.join(panel_master_path, folder_name)
    
            if os.path.isdir(full_folder_path):
                for file_name in os.listdir(full_folder_path):
                    if file_name.endswith("Final.json"):
                        source_file_path = os.path.join(full_folder_path, file_name)
                        destination_file_path = os.path.join(intermediate_path, file_name)
    
                        shutil.copyfile(source_file_path, destination_file_path)
                        print(f"File {file_name} copied to {intermediate_path}")
    
        # Merge copied JSON files into a single JSON file
        merged_data = []
    
        for file in os.listdir(intermediate_path):
            if file.endswith('Final.json'):
                with open(os.path.join(intermediate_path, file), 'r', encoding='utf-8') as json_file:
                    data = json.load(json_file)
                    for item in data['videosScriptsInfo']:
                        merged_data.append(item)
    
        result_dict = {"videosScriptsInfo": merged_data}
    
        # Save the merged data as a new JSON file
        with open(post_request_json, 'w', encoding='utf-8') as output_file:
            json.dump(result_dict, output_file, ensure_ascii=False, indent=4)

def flow_debug(panel_master_path):
    """
    Check a folder for Excel (.xlsx) and Word (.docx) files for specific rules.

    Parameters:
    - panel_master_path (str): The path to the main folder containing the files.
    - sheet_names (list): A list of sheet names to check in Excel files.
    - expected_columns (list): A list of expected column names for the first sheet in Excel files.
    - expected_columns_sheet1 (list): A list of expected column names for the second sheet in Excel files.

    Returns:
    - xlsx_not_following_rules (list): List of issues found in Excel files.
    - docx_not_following_rules (list): List of issues found in Word files.
    """

    sheet_names = ['0', '1']
    expected_columns = ["Course Name", "C", "V", "P", "L", "S", "Alt", "Question Level", "Question",
                        "Question Type", "Choice 1", "Choice 2", "Choice 3", "Choice 4", "Choice 5",
                        "Choice 6", "Right Answer"]
    expected_columns_sheet1 = ["Index", "Item (English)", "Item (Arabic)"]
    
    xlsx_not_following_rules = []
    docx_not_following_rules = []

    for root, dirs, files in os.walk(panel_master_path):
        for name in files:
            if name.endswith(".xlsx"):
                filepath = os.path.join(root, name)
                workbook = openpyxl.load_workbook(filepath)

                for sheet_name in sheet_names:
                    if sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]

                        # Check columns for the first sheet
                        if sheet_name == '0':
                            for col in expected_columns:
                                if col not in [sheet.cell(1, column_index).value for column_index in range(1, sheet.max_column + 1)]:
                                    xlsx_not_following_rules.append(name + f" in {sheet_name}: Column {col} is missing.")

                            # Check specific conditions for each row
                            for row in range(2, sheet.max_row + 1):
                                if not isinstance(sheet.cell(row, 7).value, (int, float)):
                                    xlsx_not_following_rules.append(name + f" in {sheet_name}: Cell in Alt column is not purely numerical in row {row}.")

                                for col in [3, 4, 5, 6, 8, 16]:
                                    if not (isinstance(sheet.cell(row, col).value, str) and sheet.cell(row, col).value.startswith('[') and sheet.cell(row, col).value.endswith(']')):
                                        xlsx_not_following_rules.append(name + f" in {sheet_name}: Cell in column {expected_columns[col - 1]} is not encapsulated properly in row {row}.")

                                if sheet.cell(row, 10).value == "true&false":
                                    choices = [sheet.cell(row, i).value for i in range(11, 17)]
                                    if not all(choices[:2]) and not all(choices[2:]):
                                        xlsx_not_following_rules.append(name + f" in {sheet_name}: True/false question has more than 2 choices in row {row}.")

                        # Check columns for the second sheet
                        elif sheet_name == '1':
                            for col in expected_columns_sheet1:
                                if col not in [sheet.cell(1, column_index).value for column_index in range(1, sheet.max_column + 1)]:
                                    xlsx_not_following_rules.append(name + f" in {sheet_name}: Column {col} is missing.")

            elif name.endswith(".docx"):
                doc_filepath = os.path.join(root, name)
                doc = Document(doc_filepath)
                text = " ".join([p.text for p in doc.paragraphs])
                if "[V1]" not in text:
                    docx_not_following_rules.append(name + ": Mark [V1] is missing.")

    print("List of .xlsx files not following the requirements:")
    for item_xlsx in xlsx_not_following_rules:
        print(item_xlsx)
    
    print("\nList of .docx files not following the requirements:")
    for item_docx in docx_not_following_rules:
        print(item_docx)

    return item_xlsx, item_docx
